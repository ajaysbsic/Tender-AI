import pdfplumber
from docx import Document
from typing import List, Dict, Generator, Tuple, Optional
import re
import logging
from langdetect import detect, DetectorFactory, LangDetectException

# Seed for consistent language detection
DetectorFactory.seed = 0
logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse PDF and DOCX documents with streaming, cleanup, and language detection"""
    
    # Heuristic patterns for header/footer detection
    HEADER_FOOTER_PATTERNS = [
        re.compile(r"^(Page\s+\d+|Page\s+-\s*\d+)$", re.IGNORECASE),
        re.compile(r"^(RFP|Tender|Confidential|Draft)\s+\d{4}-\d{2}-\d{2}$", re.IGNORECASE),
        re.compile(r"^(Confidential|Internal|Proprietary|Restricted)$", re.IGNORECASE),
        re.compile(r"^\d{1,2}/\d{1,2}/\d{2,4}$"),  # Dates like 01/22/2026
        re.compile(r"^http[s]?://", re.IGNORECASE),  # URLs
        re.compile(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}$"),  # Emails
        re.compile(r"^(www\.|[\w-]+\.\w+)$", re.IGNORECASE),  # Websites
    ]
    
    # Section header patterns
    SECTION_PATTERN = re.compile(r"^(\d{1,2}\.?\d*\.?\d*\s+[A-Z]|[A-Z][A-Z\s]{5,}|^#{1,6}\s)", re.MULTILINE)
    
    # Minimum text length for a page to be considered valid
    MIN_PAGE_TEXT_LENGTH = 50
    
    # Minimum text length for language detection
    MIN_LANG_DETECT_LENGTH = 100
    
    @staticmethod
    def stream_pdf_pages(file_path: str, batch_size: int = 1) -> Generator[Dict, None, None]:
        """
        Stream PDF pages one by one (or in batches) to avoid full memory load.
        Useful for handling 1000+ page documents.
        
        Args:
            file_path: Path to PDF file
            batch_size: Number of pages to process together (default 1 for streaming)
            
        Yields:
            Dictionary with page number, text, and metadata
        """
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"Opening PDF with {total_pages} pages")
                
                for i, page in enumerate(pdf.pages):
                    try:
                        text = page.extract_text() or ""
                        
                        # Skip mostly empty pages
                        if len(text.strip()) < DocumentParser.MIN_PAGE_TEXT_LENGTH:
                            logger.debug(f"Skipping empty page {i + 1}")
                            continue
                        
                        # Clean page text
                        cleaned_text = DocumentParser.remove_headers_footers(text)
                        
                        # Detect language
                        language = DocumentParser.detect_language(cleaned_text)
                        
                        yield {
                            "page": i + 1,
                            "total_pages": total_pages,
                            "text": cleaned_text,
                            "language": language,
                            "raw_length": len(text),
                            "cleaned_length": len(cleaned_text)
                        }
                        
                    except Exception as e:
                        logger.warning(f"Error processing page {i + 1}: {str(e)}")
                        continue
                        
        except Exception as e:
            logger.error(f"Error opening PDF {file_path}: {str(e)}")
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def stream_docx_paragraphs(file_path: str) -> Generator[Dict, None, None]:
        """
        Stream DOCX paragraphs (treating logical sections as 'pages').
        Useful for avoiding full memory load on large DOCX files.
        
        Yields:
            Dictionary with section number, text, and metadata
        """
        try:
            doc = Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            logger.info(f"Opening DOCX with {total_paragraphs} paragraphs")
            
            current_section = []
            section_num = 0
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                
                # Group paragraphs into logical sections (when hitting headers)
                if text and (text[0].isupper() or text.startswith("#")):
                    if current_section:
                        section_num += 1
                        section_text = "\n".join(current_section)
                        language = DocumentParser.detect_language(section_text)
                        
                        yield {
                            "section": section_num,
                            "total_sections": total_paragraphs,
                            "text": section_text,
                            "language": language,
                            "paragraph_count": len(current_section)
                        }
                        current_section = []
                
                if text:
                    current_section.append(text)
            
            # Yield final section
            if current_section:
                section_num += 1
                section_text = "\n".join(current_section)
                language = DocumentParser.detect_language(section_text)
                
                yield {
                    "section": section_num,
                    "total_sections": section_num,
                    "text": section_text,
                    "language": language,
                    "paragraph_count": len(current_section)
                }
                
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    @staticmethod
    def remove_headers_footers(text: str) -> str:
        """
        Remove headers and footers from page text using heuristics.
        
        Removes lines that match common header/footer patterns.
        """
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            stripped = line.strip()
            
            # Skip empty lines
            if not stripped:
                cleaned_lines.append("")
                continue
            
            # Check against header/footer patterns
            is_header_footer = False
            for pattern in DocumentParser.HEADER_FOOTER_PATTERNS:
                if pattern.match(stripped):
                    is_header_footer = True
                    logger.debug(f"Removing header/footer: {stripped[:50]}")
                    break
            
            # Also remove very short lines that look like page numbers
            if not is_header_footer and len(stripped) < 10:
                # Check if line is just digits or very short identifiers
                if stripped.isdigit() or (len(stripped) <= 5 and not stripped.isalpha()):
                    is_header_footer = True
            
            if not is_header_footer:
                cleaned_lines.append(line)
        
        # Remove trailing empty lines
        while cleaned_lines and not cleaned_lines[-1].strip():
            cleaned_lines.pop()
        
        return '\n'.join(cleaned_lines).strip()
    
    @staticmethod
    def detect_language(text: str) -> str:
        """
        Detect language of text using langdetect.
        
        Returns:
            ISO 639-1 language code (e.g., 'en', 'fr', 'es')
            Returns 'unknown' if detection fails or text too short
        """
        if len(text.strip()) < DocumentParser.MIN_LANG_DETECT_LENGTH:
            return "unknown"
        
        try:
            language = detect(text)
            return language
        except LangDetectException:
            logger.debug(f"Could not detect language for text: {text[:50]}")
            return "unknown"
        except Exception as e:
            logger.warning(f"Language detection error: {str(e)}")
            return "unknown"
    
    @staticmethod
    def extract_pages(file_path: str) -> List[Dict]:
        """
        Extract all pages from PDF (non-streaming, backward compatible).
        Good for documents < 500 pages.
        """
        pages = []
        for page_data in DocumentParser.stream_pdf_pages(file_path):
            pages.append(page_data)
        return pages
    
    @staticmethod
    def extract_docx_content(file_path: str) -> List[Dict]:
        """
        Extract all sections from DOCX (non-streaming, backward compatible).
        """
        sections = []
        for section_data in DocumentParser.stream_docx_paragraphs(file_path):
            sections.append(section_data)
        return sections
    
    @staticmethod
    def parse_document(file_path: str, streaming: bool = False) -> \
            Generator[Dict, None, None] | List[Dict]:
        """
        Parse document (PDF or DOCX).
        
        Args:
            file_path: Path to document
            streaming: If True, return generator; if False, return list
            
        Returns:
            Generator or List of page/section dictionaries
        """
        if file_path.lower().endswith('.pdf'):
            if streaming:
                return DocumentParser.stream_pdf_pages(file_path)
            else:
                return DocumentParser.extract_pages(file_path)
        elif file_path.lower().endswith('.docx'):
            if streaming:
                return DocumentParser.stream_docx_paragraphs(file_path)
            else:
                return DocumentParser.extract_docx_content(file_path)
        else:
            raise ValueError("Unsupported file format. Supported: PDF, DOCX")
    
    @staticmethod
    def split_by_sections(pages: List[Dict]) -> Dict[str, str]:
        """Detect and split document by sections"""
        sections = {}
        current_section = "GENERAL"
        
        for page in pages:
            text = page.get('text', '')
            for line in text.split('\n'):
                line_stripped = line.strip()
                if DocumentParser.SECTION_PATTERN.match(line_stripped):
                    current_section = line_stripped[:100]
                    if current_section not in sections:
                        sections[current_section] = []
                
                if current_section not in sections:
                    sections[current_section] = []
                sections[current_section].append(line)
        
        return {k: '\n'.join(v) for k, v in sections.items()}
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 150) -> List[str]:
        """
        Split text into overlapping chunks for AI processing.
        Smart chunking: avoids breaking in middle of sentences.
        
        Args:
            text: Text to chunk
            chunk_size: Approximate chunk size in characters
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # If not at end of text, try to break at sentence boundary
            if end < len(text):
                # Look for last sentence boundary within chunk
                last_period = text.rfind('. ', start, end)
                last_newline = text.rfind('\n', start, end)
                break_point = max(last_period, last_newline)
                
                if break_point > start:
                    end = break_point + 1
            
            chunk = text[start:end].strip()
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks
    
    @staticmethod
    def get_document_metadata(file_path: str) -> Dict:
        """
        Get metadata about document without full parsing.
        Useful for preprocessing checks.
        
        Returns:
            Dictionary with document info (page count, language, etc.)
        """
        try:
            if file_path.lower().endswith('.pdf'):
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    
                    # Sample first page for language detection
                    if page_count > 0:
                        first_page_text = pdf.pages[0].extract_text() or ""
                        language = DocumentParser.detect_language(first_page_text)
                    else:
                        language = "unknown"
                    
                    return {
                        "file_type": "pdf",
                        "page_count": page_count,
                        "estimated_language": language,
                        "file_path": file_path
                    }
            
            elif file_path.lower().endswith('.docx'):
                doc = Document(file_path)
                para_count = len(doc.paragraphs)
                
                # Sample first few paragraphs
                sample_text = "\n".join([p.text for p in doc.paragraphs[:5]])
                language = DocumentParser.detect_language(sample_text)
                
                return {
                    "file_type": "docx",
                    "paragraph_count": para_count,
                    "estimated_language": language,
                    "file_path": file_path
                }
            
            else:
                raise ValueError("Unsupported file format")
                
        except Exception as e:
            logger.error(f"Error getting metadata: {str(e)}")
            return {"error": str(e)}
